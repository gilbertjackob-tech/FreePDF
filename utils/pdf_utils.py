"""Minimal PDF utility shims used by the local dev server.

These functions provide simple, best-effort implementations so the
frontend can be exercised. They are intentionally limited and must be
replaced with robust implementations for production.
"""
import os
import shutil
import io
import re
from collections import Counter
from typing import List, Tuple
try:
    from PIL import Image
except Exception:
    Image = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None


def _sanitize_for_xml(s):
    """Remove control characters that are invalid in XML/text nodes used by python-docx.
    Keeps common whitespace but strips NULLs and C0 control codes that cause lxml errors.
    """
    if s is None:
        return ''
    if not isinstance(s, str):
        try:
            s = str(s)
        except Exception:
            return ''
    # remove C0 control chars except TAB(0x09), LF(0x0A), CR(0x0D)
    return re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', s)

def _kmeans_1d(values, k=2, max_iters=50):
    """Simple 1D KMeans implementation returning cluster centers and labels."""
    if not values:
        return [], []
    vals = list(values)
    # init centers spread across quantiles
    vals_sorted = sorted(vals)
    centers = []
    n = len(vals_sorted)
    for i in range(k):
        idx = int((i + 0.5) * n / k)
        centers.append(vals_sorted[min(idx, n-1)])
    labels = [0] * len(vals)
    for _ in range(max_iters):
        changed = False
        # assign
        for i, v in enumerate(vals):
            best = min(range(k), key=lambda c: abs(v - centers[c]))
            if labels[i] != best:
                labels[i] = best
                changed = True
        # update
        new_centers = []
        for c in range(k):
            members = [vals[i] for i in range(len(vals)) if labels[i] == c]
            if members:
                new_centers.append(sum(members) / len(members))
            else:
                new_centers.append(centers[c])
        centers = new_centers
        if not changed:
            break
    return centers, labels


def preprocess_image_for_ocr(pil_image):
    """Deskew and binarize a PIL image to improve OCR results.
    Uses OpenCV if available; otherwise attempts pytesseract OSD for orientation
    and converts to grayscale + adaptive threshold.
    """
    try:
        import cv2
        import numpy as np
        img = pil_image.convert('RGB')
        arr = np.array(img)
        gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
        # deskew using moments
        coords = cv2.findNonZero(cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1])
        if coords is not None and len(coords) > 0:
            rect = cv2.minAreaRect(coords)
            angle = rect[-1]
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            (h, w) = gray.shape
            center = (w // 2, h // 2)
            M = cv2.getRotationMatrix2D(center, angle, 1.0)
            gray = cv2.warpAffine(gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
        # adaptive threshold
        th = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 15, 10)
        return Image.fromarray(th)
    except Exception:
        # fallback: try pytesseract orientation then simple binarize
        try:
            import pytesseract
            osd = pytesseract.image_to_osd(pil_image)
            # parse rotation
            m = re.search(r'Rotate:\s*(\d+)', osd)
            if m:
                angle = int(m.group(1))
                if angle:
                    pil_image = pil_image.rotate(360 - angle, expand=True)
        except Exception:
            pass
        # grayscale and simple threshold
        try:
            gray = pil_image.convert('L')
            bw = gray.point(lambda x: 0 if x < 180 else 255, '1')
            return bw.convert('L')
        except Exception:
            return pil_image

def extract_text_from_pdf(in_path: str, ocr_enabled: bool = True, ocr_lang: str = 'eng', forced_ocr_pages: List[int] = None) -> Tuple[List[str], str, List[str]]:
    """
    Try to extract text for each page of `in_path` using a cascade of methods:
    1. PyPDF2 text extraction (fast, preserves layout when possible)
    2. PyMuPDF (fitz) page.get_text('text')
    3. OCR via pytesseract (if enabled and available) by rasterizing pages

    Returns a tuple (list_of_page_texts, backend_name) where backend_name is
    one of: 'pypdf2', 'fitz', 'ocr', or 'none'. The list may contain empty
    strings when a page produced no text.
    """
    text_pages: List[str] = []
    per_page_backends: List[str] = []

    forced = set((forced_ocr_pages or []))
    # 1) PyPDF2
    try:
        if 'PyPDF2' in globals() and PyPDF2 is not None:
            reader = PyPDF2.PdfReader(in_path)
            for idx, p in enumerate(reader.pages):
                try:
                    t = p.extract_text() or ''
                except Exception:
                    t = ''
                # if forced OCR for this page, skip text extraction and mark for OCR
                if (idx + 1) in forced:
                    text_pages.append('')
                    per_page_backends.append('force-ocr')
                else:
                    text_pages.append(t)
                    per_page_backends.append('pypdf2' if (t or '').strip() else 'none')
            if any(p.strip() for p in text_pages):
                # for pages marked force-ocr we will handle OCR below if necessary
                pass
    except Exception:
        text_pages = []
        per_page_backends = []

    # 2) PyMuPDF (fitz)
    try:
        import fitz
        try:
            doc = fitz.open(in_path)
            for i in range(len(doc)):
                try:
                    page = doc.load_page(i)
                    t = page.get_text('text') or ''
                except Exception:
                    t = ''
                if (i + 1) in forced:
                    text_pages.append('')
                    per_page_backends.append('force-ocr')
                else:
                    text_pages.append(t)
                    per_page_backends.append('fitz' if (t or '').strip() else 'none')
            doc.close()
            # if any text exists (except forced-ocr placeholders), we can return and later OCR forced pages
            pass
        except Exception:
            text_pages = []
    except Exception:
        # fitz not installed
        pass

    # 3) OCR (pytesseract) fallback: rasterize pages and OCR them
    if not ocr_enabled:
        return text_pages, 'none', per_page_backends or ['none'] * max(1, len(text_pages))

    try:
        import pytesseract
    except Exception:
        pytesseract = None

    # Need PIL for image handling
    try:
        from PIL import Image
    except Exception:
        Image = None

    if pytesseract is None or Image is None:
        return text_pages, 'none', per_page_backends or ['none'] * max(1, len(text_pages))

    images = []
    # Prefer rendering with fitz if available (for OCR rendering)
    try:
        import fitz
        doc = fitz.open(in_path)
        for i in range(len(doc)):
            try:
                page = doc.load_page(i)
                pix = page.get_pixmap(alpha=False)
                img_bytes = pix.tobytes('png')
                img = Image.open(io.BytesIO(img_bytes))
                images.append(img)
            except Exception:
                images.append(None)
        doc.close()
    except Exception:
        # fallback to pdf2image conversion
        try:
            from pdf2image import convert_from_path
            try:
                images = convert_from_path(in_path, dpi=200)
            except Exception:
                images = []
        except Exception:
            images = []

    if not images:
        return text_pages, 'none'

    ocr_pages: List[str] = []
    for idx, img in enumerate(images):
        if img is None:
            ocr_pages.append('')
            continue
        # If this page wasn't empty and not forced, skip OCR (unless forced)
        page_index = idx + 1
        if per_page_backends and per_page_backends[idx] not in ('none', 'force-ocr') and page_index not in forced:
            # keep existing extracted text if present
            ocr_pages.append(text_pages[idx] if idx < len(text_pages) else '')
            continue
        try:
            img_proc = preprocess_image_for_ocr(img)
            text = pytesseract.image_to_string(img_proc, lang=ocr_lang)
        except Exception:
            try:
                text = pytesseract.image_to_string(img.convert('RGB'), lang=ocr_lang)
            except Exception:
                text = ''
        ocr_pages.append(text or '')

    if any(p.strip() for p in ocr_pages):
        per_page_backends = ['ocr' if (p or '').strip() else 'none' for p in ocr_pages]
        return ocr_pages, 'ocr', per_page_backends

    return ocr_pages, 'none', ['none'] * max(1, len(ocr_pages))


def format_extracted_texts(text_pages: List[str]) -> List[str]:
    """
    Apply lightweight formatting to extracted page texts to improve DOCX quality.
    Heuristics applied:
    - Strip identical header/footer lines that repeat across most pages
    - Remove simple page-number lines ("Page 1", "1")
    - Remove tiny boilerplate lines that consist only of digits or very short tokens
    - Collapse excessive blank lines

    Returns cleaned list of page strings in the same order.
    """
    if not text_pages:
        return text_pages

    # Split pages into lines
    pages_lines = [ [l.rstrip() for l in (p or '').splitlines()] for p in text_pages ]
    page_count = len(pages_lines)

    # Find common first-line header and last-line footer candidates
    first_lines = [ (lines[0] if lines else '').strip() for lines in pages_lines ]
    last_lines = [ (lines[-1] if lines else '').strip() for lines in pages_lines ]

    def most_common_candidate(lines_list):
        cnt = Counter([l for l in lines_list if l])
        if not cnt:
            return ''
        val, freq = cnt.most_common(1)[0]
        # if it appears on more than half the pages, treat as common header/footer
        if freq >= max(2, page_count // 2):
            return val
        return ''

    header = most_common_candidate(first_lines)
    footer = most_common_candidate(last_lines)

    cleaned_pages = []
    page_num_re = re.compile(r'^\s*(Page\s+\d+|\d+)\s*$', re.IGNORECASE)

    for lines in pages_lines:
        # Remove header/footer if present
        if header and lines and lines[0].strip() == header:
            lines = lines[1:]
        if footer and lines and lines[-1].strip() == footer:
            lines = lines[:-1]

        # Filter out obvious page numbers and tiny numeric lines
        new_lines = []
        for ln in lines:
            s = ln.strip()
            if not s:
                new_lines.append('')
                continue
            if page_num_re.match(s):
                continue
            # drop lines that are too short and contain only digits/symbols
            if len(s) <= 3 and re.fullmatch(r'[\d\W]+', s):
                continue
            new_lines.append(ln)

        # Collapse multiple blank lines
        collapsed = []
        blank = False
        for ln in new_lines:
            if not ln.strip():
                if not blank:
                    collapsed.append('')
                blank = True
            else:
                collapsed.append(ln)
                blank = False

        cleaned_pages.append('\n'.join(collapsed).strip())

    return cleaned_pages


def split_pdf_custom(in_path: str, ranges: str, out_path: str):
    """Very small splitter: if PyPDF2 available, extract specified pages.
    ranges is a comma-separated list like '1,3-5'. Output is a single PDF
    (for simplicity) stored at out_path.
    """
    if PyPDF2 is None:
        # fallback: copy input to output so UI can download something
        shutil.copyfile(in_path, out_path)
        return
    reader = PyPDF2.PdfReader(in_path)
    writer = PyPDF2.PdfWriter()
    pages_to_keep = []
    for part in (ranges or '').split(','):
        part = part.strip()
        if not part:
            continue
        if '-' in part:
            a,b = part.split('-',1)
            pages_to_keep.extend(range(int(a)-1, int(b)))
        else:
            pages_to_keep.append(int(part)-1)
    pages_to_keep = [p for p in pages_to_keep if 0 <= p < len(reader.pages)]
    if not pages_to_keep:
        # nothing selected -> copy original
        shutil.copyfile(in_path, out_path)
        return
    for p in pages_to_keep:
        writer.add_page(reader.pages[p])
    with open(out_path, 'wb') as fh:
        writer.write(fh)

def merge_pdfs(paths: List[str], out_path: str):
    import time, logging
    logger = logging.getLogger('pdf_utils')
    start = time.time()
    # Try pikepdf first (fast, stream-optimized)
    try:
        import pikepdf
        logger.info('merge_pdfs: using pikepdf for %d files', len(paths))
        with pikepdf.Pdf.new() as out_pdf:
            for p in paths:
                try:
                    with pikepdf.Pdf.open(p) as src:
                        out_pdf.pages.extend(src.pages)
                except Exception:
                    # skip problematic file but continue merging
                    logger.exception('pikepdf failed for %s', p)
            out_pdf.save(out_path)
        logger.info('merge_pdfs: pikepdf merge completed in %.2fs', time.time() - start)
        return
    except Exception:
        logger.debug('pikepdf not available or failed, falling back')

    # Next try PyPDF2's PdfMerger if available (more efficient than manual page loop)
    try:
        from PyPDF2 import PdfMerger
        logger.info('merge_pdfs: using PyPDF2.PdfMerger for %d files', len(paths))
        merger = PdfMerger()
        for p in paths:
            try:
                merger.append(p)
            except Exception:
                logger.exception('PdfMerger append failed for %s', p)
        merger.write(out_path)
        merger.close()
        logger.info('merge_pdfs: PdfMerger completed in %.2fs', time.time() - start)
        return
    except Exception:
        logger.debug('PdfMerger not available or failed, falling back to page-by-page')

    # Fallback: page-by-page add (works with older PyPDF2 or when nothing else available)
    if PyPDF2 is None:
        # fallback: copy first file
        if paths:
            shutil.copyfile(paths[0], out_path)
        else:
            open(out_path, 'wb').close()
        logger.info('merge_pdfs: fallback copy completed in %.2fs', time.time() - start)
        return
    writer = PyPDF2.PdfWriter()
    for p in paths:
        try:
            r = PyPDF2.PdfReader(p)
            for pg in r.pages:
                writer.add_page(pg)
        except Exception:
            logger.exception('page-by-page merge failed for %s', p)
    with open(out_path, 'wb') as fh:
        writer.write(fh)
    logger.info('merge_pdfs: page-by-page merge completed in %.2fs', time.time() - start)

    # Basic validation: ensure output looks like a PDF (helps detect silent failures)
    try:
        with open(out_path, 'rb') as fh:
            sig = fh.read(4)
            if not sig.startswith(b'%PDF'):
                raise RuntimeError('merged file does not start with %PDF header')
    except Exception:
        # If validation fails, remove the bad file to avoid returning a corrupt artifact
        try:
            os.remove(out_path)
        except Exception:
            pass
        raise


def images_to_pdf_util(image_paths: List[str], out_path: str, margin_mm: float = 10):
    if Image is None:
        open(out_path, 'wb').close()
        return

    imgs = []
    for p in image_paths:
        try:
            img = Image.open(p).convert('RGB')
            imgs.append(img)
        except Exception:
            continue

    if not imgs:
        open(out_path, 'wb').close()
        return

    # A4 at 300 DPI
    DPI = 300
    A4_W_IN, A4_H_IN = 8.27, 11.69  # inches
    a4_w = int(A4_W_IN * DPI)
    a4_h = int(A4_H_IN * DPI)

    # margin in pixels
    margin_px = int((margin_mm / 25.4) * DPI)  # mm -> inches -> pixels

    pdf_pages = []
    for img in imgs:
        # Resize image to fit inside A4 with margin, keeping aspect ratio
        w, h = img.size
        max_w = a4_w - 2 * margin_px
        max_h = a4_h - 2 * margin_px
        scale = min(max_w / w, max_h / h)
        new_w, new_h = int(w * scale), int(h * scale)
        img_resized = img.resize((new_w, new_h), Image.LANCZOS)

        # Paste onto A4 white background
        canvas = Image.new('RGB', (a4_w, a4_h), (255, 255, 255))
        px = (a4_w - new_w) // 2
        py = (a4_h - new_h) // 2
        canvas.paste(img_resized, (px, py))
        pdf_pages.append(canvas)

    first, rest = pdf_pages[0], pdf_pages[1:]
    first.save(out_path, save_all=True, append_images=rest)


def pdf_to_docx(in_path: str, out_path: str, ocr_lang: str = 'eng', *,
                fragment_threshold: int = 8,
                non_black_pct: float = 0.02,
                small_image_px: int = 64*64,
                area_frac_thresh: float = 0.8,
                stdev_thresh: float = 2.5,
                mean_thresh: float = 12.0,
                enable_table_detection: bool = False,
                table_detection_mode: str = 'light'):
    """
    Simplified PDF -> DOCX conversion WITHOUT OCR.
    - Prefer PyMuPDF (fitz) structured extraction and embed images.
    - If fitz is not available, fall back to PyPDF2 text extraction (no OCR).
    Returns a diagnostics dict: {'backend': 'fitz'|'pypdf2'|'none', 'per_page': [...], 'images': [...]}
    """
    try:
        from docx import Document
    except Exception:
        try:
            open(out_path, 'wb').close()
        except Exception:
            pass
        return {'backend': 'none', 'per_page': [], 'images': []}

    # Try PyMuPDF structured extraction first (no OCR)
    try:
        import fitz
    except Exception:
        fitz = None

    if fitz:
        try:
            doc_fz = fitz.open(in_path)
            from docx.shared import Inches, Pt
            from io import BytesIO
            doc = Document()
            per_page = []
            image_counts = []
            for pi in range(len(doc_fz)):
                page = doc_fz.load_page(pi)
                d = page.get_text('dict')
                spans = []
                try:
                    for block in d.get('blocks', []):
                        if block.get('type') != 0:
                            continue
                        for line in block.get('lines', []):
                            for span in line.get('spans', []):
                                txt = span.get('text', '')
                                if not txt.strip():
                                    continue
                                bbox = span.get('bbox') or [0,0,0,0]
                                spans.append({'y0': bbox[1], 'x0': bbox[0], 'text': txt, 'size': float(span.get('size', 10)), 'font': span.get('font', ''), 'flags': span.get('flags', 0)})
                except Exception:
                    pass

                # Default paragraph grouping function
                def _write_paragraphs(spans_list):
                    last_y = None
                    para = None
                    for s in spans_list:
                        if last_y is None or abs(s['y0'] - last_y) > max(6.0, s.get('size', 10) * 0.8):
                            para = doc.add_paragraph()
                        text_safe = _sanitize_for_xml(s.get('text', ''))
                        run = para.add_run(text_safe)
                        try:
                            run.font.size = Pt(s.get('size', 10))
                            if s.get('font'):
                                run.font.name = s.get('font')
                            fname = (s.get('font') or '').lower()
                            if 'bold' in fname or 'bd' in fname or (s.get('flags', 0) & 2):
                                run.bold = True
                            if 'italic' in fname or 'oblique' in fname or (s.get('flags', 0) & 1):
                                run.italic = True
                        except Exception:
                            pass
                        last_y = s['y0']

                # table detection: two modes supported
                # - 'light' (default): original conservative heuristic
                # - 'robust' : stronger clustering-based detection (opt-in)
                if enable_table_detection and (table_detection_mode or 'light') == 'light':
                    spans_sorted = sorted(spans, key=lambda s: (s['y0'], s['x0']))
                    # bucket spans into lines by rounded y0
                    lines = []
                    cur_y = None
                    for s in spans_sorted:
                        y = float(s.get('y0', 0.0))
                        if cur_y is None or abs(y - cur_y) > max(6.0, s.get('size', 10) * 0.8):
                            lines.append([])
                            cur_y = y
                        lines[-1].append(s)

                    # compute repeated x positions across lines (rounded to 10pt grid)
                    from collections import Counter
                    x_counts = Counter()
                    for ln in lines:
                        seen = set()
                        for s in ln:
                            rx = int(round(float(s.get('x0', 0.0)) / 10.0) * 10)
                            if rx not in seen:
                                x_counts[rx] += 1
                                seen.add(rx)

                    # be more sensitive for sparser tables (allow as few as 2 repeat rows)
                    min_repeat = max(2, len(lines) // 8)
                    candidate_columns = sorted([x for x, c in x_counts.items() if c >= min_repeat])

                    is_table = False
                    if len(candidate_columns) >= 2 and len(lines) >= 3:
                        # check average coverage: many lines should have spans at these columns
                        hits = 0
                        for ln in lines:
                            rx_set = set(int(round(float(s.get('x0', 0.0)) / 10.0) * 10) for s in ln)
                            # count how many candidate columns are near spans in this line
                            for c in candidate_columns:
                                # consider a hit if a span lies within ~20pt of the column x
                                if any(abs(rx - c) <= 20 for rx in rx_set):
                                    hits += 1
                                    break
                        avg_hits = hits / max(1, len(lines))
                        if avg_hits >= max(1.0, len(candidate_columns) * 0.45):
                            is_table = True

                    if is_table:
                        # create table with rows=len(lines) and cols=len(candidate_columns)
                        cols = len(candidate_columns)
                        rows = len(lines)
                        try:
                            tbl = doc.add_table(rows=rows, cols=cols)
                            tbl.autofit = True
                            for r_idx, ln in enumerate(lines):
                                # map spans in this line into the candidate columns
                                cells_text = [''] * cols
                                for s in ln:
                                    rx = int(round(float(s.get('x0', 0.0)) / 10.0) * 10)
                                    # find nearest candidate column (assign to nearest col)
                                    best_col = None
                                    best_dist = None
                                    for ci, cval in enumerate(candidate_columns):
                                        d = abs(rx - cval)
                                        if best_dist is None or d < best_dist:
                                            best_dist = d
                                            best_col = ci
                                    # only assign if reasonably close (20pt) else append to nearest anyway
                                    if best_col is not None:
                                        if cells_text[best_col]:
                                            cells_text[best_col] += ' ' + s.get('text', '')
                                        else:
                                            cells_text[best_col] = s.get('text', '')
                                for c_idx, text in enumerate(cells_text):
                                    try:
                                        cell = tbl.rows[r_idx].cells[c_idx]
                                        # simple set text (no styling)
                                        cell.text = _sanitize_for_xml((text or '').strip())
                                    except Exception:
                                        pass
                        except Exception:
                            # fallback to paragraph fallback if table creation fails
                            is_table = False

                    if not is_table:
                        _write_paragraphs(spans_sorted)
                elif enable_table_detection and table_detection_mode == 'robust':
                    # Robust table detection: cluster x positions and require consistent column hits.
                    spans_sorted = sorted(spans, key=lambda s: (s['y0'], s['x0']))
                    # bucket spans into lines by rounded y0
                    lines = []
                    cur_y = None
                    for s in spans_sorted:
                        y = float(s.get('y0', 0.0))
                        if cur_y is None or abs(y - cur_y) > max(6.0, s.get('size', 10) * 0.8):
                            lines.append([])
                            cur_y = y
                        lines[-1].append(s)

                    # collect x positions and quantize to 5pt grid for stability
                    x_vals = [float(s.get('x0', 0.0)) for s in spans_sorted]
                    if not x_vals:
                        _write_paragraphs(spans_sorted); is_table = False
                    else:
                        qx = [int(round(x / 5.0) * 5) for x in x_vals]
                        from collections import Counter
                        qcnt = Counter(qx)
                        # candidate peaks (strong columns)
                        min_repeat = max(2, len(lines) // 6)
                        peaks = sorted([x for x, c in qcnt.items() if c >= min_repeat])

                        candidate_columns = []
                        if len(peaks) >= 2:
                            # use peaks as initial columns
                            candidate_columns = peaks
                        else:
                            # fallback: try k-means clustering on raw x positions
                            uniq = sorted(set(qx))
                            best_cols = []
                            best_score = 0.0
                            max_k = min(8, max(2, len(uniq)))
                            for k in range(2, max_k + 1):
                                centers, labels = _kmeans_1d(x_vals, k=k)
                                # compute coverage: proportion of spans within 25pt of nearest center
                                hits = 0
                                for xv in x_vals:
                                    if any(abs(xv - c) <= 25 for c in centers):
                                        hits += 1
                                score = hits / max(1, len(x_vals))
                                # penalize too many tiny clusters
                                if score > best_score:
                                    best_score = score
                                    best_cols = [int(round(c / 5.0) * 5) for c in centers]
                            candidate_columns = sorted(set(best_cols))

                        is_table = False
                        if len(candidate_columns) >= 2 and len(lines) >= 3:
                            # map lines to columns and compute hit ratio
                            hits = 0
                            for ln in lines:
                                rx_set = [int(round(float(s.get('x0', 0.0)))) for s in ln]
                                matched = 0
                                for c in candidate_columns:
                                    if any(abs(rx - c) <= 30 for rx in rx_set):
                                        matched += 1
                                if matched >= 1:
                                    hits += 1
                            avg_hits = hits / max(1, len(lines))
                            # require at least ~45% line coverage to treat as table
                            if avg_hits >= max(0.45, len(candidate_columns) * 0.15):
                                is_table = True

                        if is_table:
                            cols = len(candidate_columns)
                            rows = len(lines)
                            try:
                                tbl = doc.add_table(rows=rows, cols=cols)
                                tbl.autofit = True
                                for r_idx, ln in enumerate(lines):
                                    cells_text = [''] * cols
                                    for s in ln:
                                        rx = int(round(float(s.get('x0', 0.0))))
                                        # find nearest candidate column
                                        best_col = None
                                        best_dist = None
                                        for ci, cval in enumerate(candidate_columns):
                                            d = abs(rx - cval)
                                            if best_dist is None or d < best_dist:
                                                best_dist = d
                                                best_col = ci
                                        if best_col is not None:
                                            if cells_text[best_col]:
                                                cells_text[best_col] += ' ' + s.get('text', '')
                                            else:
                                                cells_text[best_col] = s.get('text', '')
                                    for c_idx, text in enumerate(cells_text):
                                        try:
                                            cell = tbl.rows[r_idx].cells[c_idx]
                                            cell.text = _sanitize_for_xml((text or '').strip())
                                        except Exception:
                                            pass
                            except Exception:
                                is_table = False
                        if not is_table:
                            _write_paragraphs(spans_sorted)
                elif enable_table_detection and table_detection_mode == 'aggressive':
                    # Aggressive detection: relaxed clustering + optional OpenCV image-grid detection
                    spans_sorted = sorted(spans, key=lambda s: (s['y0'], s['x0']))
                    # bucket spans into lines by rounded y0
                    lines = []
                    cur_y = None
                    for s in spans_sorted:
                        y = float(s.get('y0', 0.0))
                        if cur_y is None or abs(y - cur_y) > max(6.0, s.get('size', 10) * 0.8):
                            lines.append([])
                            cur_y = y
                        lines[-1].append(s)

                    x_vals = [float(s.get('x0', 0.0)) for s in spans_sorted]
                    candidate_columns = []
                    is_table = False
                    # first attempt: very relaxed quantization and lower min_repeat
                    if x_vals:
                        qx = [int(round(x / 5.0) * 5) for x in x_vals]
                        import collections as _collections
                        qcnt = _collections.Counter(qx)
                        min_repeat = max(2, len(lines) // 10)  # more tolerant
                        peaks = sorted([x for x, c in qcnt.items() if c >= min_repeat])
                        if len(peaks) >= 2:
                            candidate_columns = peaks
                        else:
                            # fallback: cluster with k up to 10
                            uniq = sorted(set(qx))
                            best_cols = []
                            best_score = 0.0
                            max_k = min(10, max(2, len(uniq)))
                            for k in range(2, max_k + 1):
                                centers, labels = _kmeans_1d(x_vals, k=k)
                                hits = 0
                                for xv in x_vals:
                                    if any(abs(xv - c) <= 35 for c in centers):
                                        hits += 1
                                score = hits / max(1, len(x_vals))
                                if score > best_score:
                                    best_score = score
                                    best_cols = [int(round(c / 5.0) * 5) for c in centers]
                            candidate_columns = sorted(set(best_cols))

                    # Second attempt: if OpenCV is available, render page and detect table lines
                    cv_columns = []
                    try:
                        import cv2
                        import numpy as np
                        # render higher-res pixmap for better line detection
                        try:
                            mat = fitz.Matrix(2.0, 2.0)
                            pix = page.get_pixmap(matrix=mat, alpha=False)
                            img_bytes = pix.tobytes('png')
                            from PIL import Image as PILImage
                            pil = PILImage.open(io.BytesIO(img_bytes)).convert('L')
                            arr = np.array(pil)
                            # adaptive threshold to binary
                            th = cv2.adaptiveThreshold(arr, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 15, 8)
                            # morphological ops to join table lines
                            horizontal = th.copy()
                            vertical = th.copy()
                            cols_img = horizontal.shape[1]
                            horiz_size = max(10, cols_img // 40)
                            vert_size = max(10, horizontal.shape[0] // 40)
                            horiz_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (horiz_size, 1))
                            vert_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, vert_size))
                            horizontal = cv2.morphologyEx(horizontal, cv2.MORPH_OPEN, horiz_kernel)
                            vertical = cv2.morphologyEx(vertical, cv2.MORPH_OPEN, vert_kernel)
                            # combine
                            grid = cv2.add(horizontal, vertical)
                            # sum columns to find vertical line peaks
                            col_sum = np.sum(grid > 0, axis=0)
                            # find peaks in col_sum
                            peaks_idx = []
                            thresh = max(5, int(0.1 * grid.shape[0]))
                            for i in range(1, len(col_sum)-1):
                                if col_sum[i] > thresh and col_sum[i] >= col_sum[i-1] and col_sum[i] >= col_sum[i+1]:
                                    peaks_idx.append(i)
                            # convert pixel x positions back to PDF points (approx)
                            if peaks_idx:
                                scale = pix.width / page.rect.width if page.rect.width else 1.0
                                cv_columns = [int(round(px / scale)) for px in peaks_idx]
                        except Exception:
                            # image-based detection for this page failed; continue without cv results
                            cv_columns = []
                    except Exception:
                        cv_columns = []

                    # merge column sources
                    all_columns = sorted(set((candidate_columns or []) + (cv_columns or [])))
                    # sanity checks
                    if len(all_columns) >= 2 and len(lines) >= 2:
                        # compute coverage
                        hits = 0
                        for ln in lines:
                            rx_set = [int(round(float(s.get('x0', 0.0)))) for s in ln]
                            if any(any(abs(rx - c) <= 40 for c in all_columns) for rx in rx_set):
                                hits += 1
                        avg_hits = hits / max(1, len(lines))
                        if avg_hits >= 0.25:  # be permissive
                            is_table = True

                    if is_table:
                        cols = len(all_columns)
                        rows = len(lines)
                        try:
                            tbl = doc.add_table(rows=rows, cols=cols)
                            tbl.autofit = True
                            for r_idx, ln in enumerate(lines):
                                cells_text = [''] * cols
                                for s in ln:
                                    rx = int(round(float(s.get('x0', 0.0))))
                                    # assign to nearest column
                                    best_col = None
                                    best_dist = None
                                    for ci, cval in enumerate(all_columns):
                                        d = abs(rx - cval)
                                        if best_dist is None or d < best_dist:
                                            best_dist = d
                                            best_col = ci
                                    if best_col is not None:
                                        if cells_text[best_col]:
                                            cells_text[best_col] += ' ' + s.get('text', '')
                                        else:
                                            cells_text[best_col] = s.get('text', '')
                                for c_idx, text in enumerate(cells_text):
                                    try:
                                        cell = tbl.rows[r_idx].cells[c_idx]
                                        cell.text = _sanitize_for_xml((text or '').strip())
                                    except Exception:
                                        pass
                        except Exception:
                            is_table = False
                    if not is_table:
                        _write_paragraphs(spans_sorted)
                else:
                    # table detection disabled: default paragraph grouping
                    spans_sorted = sorted(spans, key=lambda s: (s['y0'], s['x0']))
                    _write_paragraphs(spans_sorted)

                # images: prefer extracting raw image bytes from the page blocks (preserve original bytes)
                images_info = []
                try:
                    for block in d.get('blocks', []):
                        if block.get('type') == 1:
                            bbox = block.get('bbox', [0, 0, 0, 0])
                            # block may contain an 'image' dict with xref
                            im = block.get('image') or {}
                            xref = None
                            if isinstance(im, dict):
                                xref = im.get('xref') or im.get('xref')
                            if not xref:
                                xref = block.get('xref') or (im if isinstance(im, int) else None)
                            if xref:
                                images_info.append({'xref': xref, 'bbox': bbox})
                except Exception:
                    images_info = []

                # fallback to page.get_images if no block-level images found
                if not images_info:
                    try:
                        imgs = page.get_images(full=True) or []
                        for img_meta in imgs:
                            xref = img_meta[0]
                            images_info.append({'xref': xref, 'bbox': None})
                    except Exception:
                        images_info = []
                # If there are many small image fragments, it's likely the PDF stores
                # a single visual as many slices; render the whole page once instead
                # to avoid inserting many thin/cropped strips.
                if len(images_info) > int(fragment_threshold):
                    try:
                        pix = page.get_pixmap(alpha=False)
                        page_img_bytes = pix.tobytes('png')
                        # insert the full-page image scaled to page width
                        try:
                            doc.add_picture(BytesIO(page_img_bytes), width=Inches(float(page.rect.width) / 72.0))
                            image_counts.append(1)
                        except Exception:
                            # fallback to default insert
                            doc.add_picture(BytesIO(page_img_bytes))
                            image_counts.append(1)
                    except Exception:
                        # fall back to per-image insertion
                        pass
                    # skip individual fragments insertion
                    if pi != len(doc_fz) - 1:
                        try:
                            doc.add_page_break()
                        except Exception:
                            pass
                    per_page.append('fitz' if spans_sorted else 'none')
                    continue

                kept_images = 0
                for im_info in images_info:
                    xref = im_info.get('xref')
                    bbox = im_info.get('bbox')
                    img_bytes = None
                    try:
                        # try to extract original image bytes (keeps original encoding)
                        try:
                            img_dict = doc_fz.extract_image(xref)
                            if img_dict and 'image' in img_dict:
                                img_bytes = img_dict['image']
                        except Exception:
                            img_bytes = None
                        if not img_bytes:
                            # fallback to Pixmap -> PNG bytes
                            pix = fitz.Pixmap(doc_fz, xref)
                            if pix.n < 5:
                                img_bytes = pix.tobytes('png')
                            else:
                                pix = fitz.Pixmap(fitz.csRGB, pix)
                                img_bytes = pix.tobytes('png')
                    except Exception:
                        img_bytes = None
                    if not img_bytes:
                        continue
                    # Inspect image content to skip near-solid-black or tiny images and to get pixel size
                    im_test = None
                    try:
                        from PIL import Image as PILImage
                        im_test = PILImage.open(BytesIO(img_bytes))
                        # convert to grayscale and downsize to speed up analysis
                        try:
                            im_small = im_test.convert('L').resize((64, 64), PILImage.Resampling.BILINEAR)
                        except Exception:
                            im_small = im_test.convert('L').resize((64, 64))
                        pixels = list(im_small.getdata())
                        if pixels:
                            # count pixels brighter than a tiny threshold
                            non_black = sum(1 for p in pixels if p > 10)
                            pct_non_black = non_black / len(pixels)
                        else:
                            pct_non_black = 0.0
                        # determine page area and bbox coverage fraction
                        page_area = float(page.rect.width * page.rect.height) if getattr(page, 'rect', None) else None
                        bbox_area = None
                        if bbox and len(bbox) >= 4:
                            try:
                                bbox_w = max(0.0, float(bbox[2]) - float(bbox[0]))
                                bbox_h = max(0.0, float(bbox[3]) - float(bbox[1]))
                                bbox_area = bbox_w * bbox_h
                            except Exception:
                                bbox_area = None
                        area_frac = (bbox_area / page_area) if (bbox_area and page_area and page_area > 0) else 0.0
                        # additional check: skip images that are near-solid dark color (low variance)
                        try:
                            import statistics
                            vals = pixels
                            mean = sum(vals) / len(vals) if vals else 0.0
                            stdev = statistics.pstdev(vals) if vals else 0.0
                        except Exception:
                            mean = 0.0
                            stdev = 0.0

                        if pct_non_black < float(non_black_pct) or ((im_test.size[0] * im_test.size[1] < int(small_image_px)) and pct_non_black < max(0.01, float(non_black_pct)*2.5)) or (area_frac > float(area_frac_thresh) and pct_non_black < max(0.01, float(non_black_pct)*2.5)) or (stdev < float(stdev_thresh) and mean < float(mean_thresh)):
                            # skip inserting this image (near-solid dark or too small/empty)
                            continue
                    except Exception:
                        # if PIL analysis fails, proceed with insertion
                        im_test = None

                    # size image based on DPI and pixel dimensions when possible to preserve native display size
                    try:
                        if im_test is not None:
                            try:
                                dpi = None
                                if hasattr(im_test, 'info') and isinstance(im_test.info, dict):
                                    dpi = im_test.info.get('dpi') or im_test.info.get('resolution')
                                if isinstance(dpi, tuple):
                                    dpi = dpi[0]
                                if not dpi:
                                    dpi = 72.0
                            except Exception:
                                dpi = 72.0
                            px_w, px_h = im_test.size
                            # compute width in inches from pixels/dpi
                            width_in = float(px_w) / float(dpi or 72.0)
                            # if bbox provided, avoid huge/inaccurate sizes: cap to bbox width if available
                            if bbox and len(bbox) >= 4 and bbox[2] > bbox[0]:
                                bbox_width_in = (float(bbox[2]) - float(bbox[0])) / 72.0
                                use_width_in = min(width_in, max(0.01, bbox_width_in))
                            else:
                                use_width_in = width_in
                            doc.add_picture(BytesIO(img_bytes), width=Inches(use_width_in))
                        else:
                            # fallback sizing: bbox or 90% page width
                            if bbox and len(bbox) >= 4 and bbox[2] > bbox[0]:
                                width_pts = bbox[2] - bbox[0]
                                width_in = float(width_pts) / 72.0
                                doc.add_picture(BytesIO(img_bytes), width=Inches(width_in))
                            else:
                                width_pts = page.rect.width
                                width_in = float(width_pts) / 72.0 * 0.9
                                doc.add_picture(BytesIO(img_bytes), width=Inches(width_in))
                    except Exception:
                        try:
                            doc.add_picture(BytesIO(img_bytes))
                        except Exception:
                            pass
                    kept_images += 1

                if pi != len(doc_fz) - 1:
                    try:
                        doc.add_page_break()
                    except Exception:
                        pass
                per_page.append('fitz' if spans_sorted else 'none')

                image_counts.append(kept_images)

            try:
                doc.save(out_path)
            except Exception:
                try:
                    open(out_path, 'wb').close()
                except Exception:
                    pass
            try:
                doc_fz.close()
            except Exception:
                pass
            return {'backend': 'fitz', 'per_page': per_page, 'images': image_counts}
        except Exception:
            # Log the traceback to a file in output/ so we can inspect why fitz branch failed,
            # then fall through to PyPDF2 fallback.
            try:
                import traceback
                out_dir = os.path.join(os.path.dirname(__file__), '..', 'output')
                os.makedirs(out_dir, exist_ok=True)
                basename = os.path.basename(in_path)
                trace_path = os.path.join(out_dir, f'fitz_exception_{basename}.txt')
                with open(trace_path, 'w', encoding='utf-8') as tf:
                    tf.write('FITZ BRANCH EXCEPTION FOR: ' + repr(in_path) + '\n')
                    tf.write(traceback.format_exc())
            except Exception:
                pass
            try:
                doc_fz.close()
            except Exception:
                pass

    # Fallback: use PyPDF2 text extraction only (no OCR)
    try:
        if 'PyPDF2' in globals() and PyPDF2 is not None:
            reader = PyPDF2.PdfReader(in_path)
            from docx import Document as DocxDocument
            doc = DocxDocument()
            per_page = []
            for i, p in enumerate(reader.pages, start=1):
                try:
                    text = p.extract_text() or ''
                except Exception:
                    text = ''
                if not text.strip():
                    per_page.append('none')
                    doc.add_paragraph(f'Page {i}: (no selectable text)')
                else:
                    per_page.append('pypdf2')
                    for part in [p.strip() for p in text.split('\n\n') if p.strip()]:
                        doc.add_paragraph(part)
                if i != len(reader.pages):
                    try:
                        doc.add_page_break()
                    except Exception:
                        pass
            try:
                doc.save(out_path)
            except Exception:
                try:
                    open(out_path, 'wb').close()
                except Exception:
                    pass
            return {'backend': 'pypdf2', 'per_page': per_page, 'images': []}
    except Exception:
        pass

    # Last-resort empty doc
    try:
        doc = Document()
        doc.add_paragraph('Converted from PDF: ' + os.path.basename(in_path))
        doc.save(out_path)
    except Exception:
        try:
            open(out_path, 'wb').close()
        except Exception:
            pass
    return {'backend': 'none', 'per_page': [], 'images': []}

def edit_pdf(in_path: str, out_path: str, rotate: int = 0):
    # very small edit: if rotate != 0 and PyPDF2 available, rotate pages
    if PyPDF2 is None:
        shutil.copyfile(in_path, out_path); return
    reader = PyPDF2.PdfReader(in_path)
    writer = PyPDF2.PdfWriter()
    for p in reader.pages:
        try:
            if rotate:
                p.rotate_clockwise(rotate)
        except Exception:
            pass
        writer.add_page(p)
    with open(out_path, 'wb') as fh:
        writer.write(fh)

def reorder_pages(in_path: str, order: List[int], out_path: str, rotate: int = 0, delete_list: List[int] = None):
    if PyPDF2 is None:
        shutil.copyfile(in_path, out_path); return
    reader = PyPDF2.PdfReader(in_path)
    writer = PyPDF2.PdfWriter()
    total = len(reader.pages)
    delete_set = set((delete_list or []))
    for idx in order:
        if idx <= 0 or idx > total: continue
        if idx in delete_set: continue
        p = reader.pages[idx-1]
        try:
            if rotate:
                p.rotate_clockwise(rotate)
        except Exception:
            pass
        writer.add_page(p)
    with open(out_path, 'wb') as fh:
        writer.write(fh)

def generate_thumbnails(in_path: str, thumb_dir: str):
    os.makedirs(thumb_dir, exist_ok=True)
    thumbs = []
    # Prefer PyMuPDF (fitz) for page rendering
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(in_path)
        for i in range(len(doc)):
            page = doc.load_page(i)
            mat = fitz.Matrix(2, 2)  # scale for higher-res thumb
            pix = page.get_pixmap(matrix=mat, alpha=False)
            dest = os.path.join(thumb_dir, f'page{i+1}.jpg')
            pix.save(dest)
            thumbs.append(dest)
        doc.close()
        if thumbs:
            return thumbs
    except Exception:
        # PyMuPDF not available or failed — try pdf2image
        pass

    # Try pdf2image (requires poppler) if PIL is available for saving
    try:
        from pdf2image import convert_from_path
        images = convert_from_path(in_path, dpi=150)
        for i, im in enumerate(images):
            dest = os.path.join(thumb_dir, f'page{i+1}.jpg')
            im.save(dest, 'JPEG')
            thumbs.append(dest)
        if thumbs:
            return thumbs
    except Exception:
        pass

    # Fallback: create placeholder images — one per PDF page if possible
    page_count = 1
    try:
        if PyPDF2:
            reader = PyPDF2.PdfReader(in_path)
            page_count = max(1, len(reader.pages))
    except Exception:
        page_count = 1

    for i in range(page_count):
        dest = os.path.join(thumb_dir, f'page{i+1}.jpg')
        if not os.path.exists(dest):
            if Image:
                img = Image.new('RGB', (300, 400), color=(240, 240, 240))
                img.save(dest, 'JPEG')
            else:
                open(dest, 'wb').close()
        thumbs.append(dest)

    return thumbs


def merge_pdfs_bytes(paths: List[str]) -> tuple:
    """Merge PDFs in-memory and return bytes of the merged PDF.
    Tries pikepdf, then PyPDF2.PdfMerger, then falls back to page-by-page into
    an in-memory buffer. This avoids writing a temporary file to disk.
    """
    import io, time, logging
    logger = logging.getLogger('pdf_utils')
    start = time.time()
    # Try pikepdf
    try:
        import pikepdf
    except Exception:
        pikepdf = None

    if pikepdf:
        try:
            logger.info('merge_pdfs_bytes: using pikepdf for %d files', len(paths))
            out_buf = io.BytesIO()
            with pikepdf.Pdf.new() as out_pdf:
                for p in paths:
                    try:
                        with pikepdf.Pdf.open(p) as src:
                            out_pdf.pages.extend(src.pages)
                    except Exception:
                        logger.exception('pikepdf failed for %s', p)
                out_pdf.save(out_buf)
            data = out_buf.getvalue()
            elapsed = time.time() - start
            # sanity check: ensure bytes look like a PDF
            if not data or not data.startswith(b'%PDF'):
                logger.warning('merge_pdfs_bytes: pikepdf produced non-PDF data, falling back')
                raise RuntimeError('pikepdf produced invalid PDF bytes')
            logger.info('merge_pdfs_bytes: pikepdf merge completed in %.2fs', elapsed)
            return data, 'pikepdf', elapsed
        except Exception:
            logger.exception('pikepdf in-memory merge failed, falling back')

    # Try PyPDF2 PdfMerger
    try:
        from PyPDF2 import PdfMerger
    except Exception:
        PdfMerger = None

    if PdfMerger:
        try:
            logger.info('merge_pdfs_bytes: using PyPDF2.PdfMerger for %d files', len(paths))
            merger = PdfMerger()
            for p in paths:
                try:
                    merger.append(p)
                except Exception:
                    logger.exception('PdfMerger append failed for %s', p)
            out_buf = io.BytesIO()
            merger.write(out_buf)
            merger.close()
            data = out_buf.getvalue()
            elapsed = time.time() - start
            if not data or not data.startswith(b'%PDF'):
                logger.warning('merge_pdfs_bytes: PdfMerger produced non-PDF data, falling back')
                raise RuntimeError('PdfMerger produced invalid PDF bytes')
            logger.info('merge_pdfs_bytes: PdfMerger completed in %.2fs', elapsed)
            return data, 'pdfmerger', elapsed
        except Exception:
            logger.exception('PdfMerger in-memory merge failed, falling back')

    # Fallback: use PyPDF2 writer page-by-page into buffer
    try:
        import PyPDF2 as _pypdf
        logger.info('merge_pdfs_bytes: using PyPDF2 page-by-page for %d files', len(paths))
        writer = _pypdf.PdfWriter()
        for p in paths:
            try:
                r = _pypdf.PdfReader(p)
                for pg in r.pages:
                    writer.add_page(pg)
            except Exception:
                logger.exception('page-by-page merge failed for %s', p)
        out_buf = io.BytesIO()
        writer.write(out_buf)
        data = out_buf.getvalue()
        elapsed = time.time() - start
        if not data or not data.startswith(b'%PDF'):
            logger.warning('merge_pdfs_bytes: page-by-page merge produced non-PDF data')
            raise RuntimeError('page-by-page merge produced invalid PDF bytes')
        logger.info('merge_pdfs_bytes: page-by-page merge completed in %.2fs', elapsed)
        return data, 'pypdf2-pageby', elapsed
    except Exception:
        logger.exception('All in-memory merge methods failed')
        return b'', 'failed', time.time() - start
